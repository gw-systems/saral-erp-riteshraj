
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT / 'LR.docx'
LOGO_PATH = ROOT / 'static' / 'images' / 'GW Logo.png'


def _set_cell_border(cell, **kwargs):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in('w:tcBorders')
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data is None:
            continue
        tag = qn(f'w:{edge}')
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        for key, value in edge_data.items():
            element.set(qn(f'w:{key}'), str(value))


def _set_table_borders(table, color='000000', size='8'):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tbl_pr.append(borders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        element = borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def _make_table_borderless(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(
                cell,
                top={'val': 'nil'},
                bottom={'val': 'nil'},
                left={'val': 'nil'},
                right={'val': 'nil'},
            )


def _set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in('w:tcMar')
    if tc_mar is None:
        tc_mar = OxmlElement('w:tcMar')
        tc_pr.append(tc_mar)
    for name, value in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = tc_mar.find(qn(f'w:{name}'))
        if node is None:
            node = OxmlElement(f'w:{name}')
            tc_mar.append(node)
        node.set(qn('w:w'), str(value))
        node.set(qn('w:type'), 'dxa')


def _set_col_widths(row, widths):
    for cell, width in zip(row.cells, widths):
        cell.width = width


def _set_paragraph_format(paragraph, before=0, after=0, line=None, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if line:
        fmt.line_spacing = Pt(line)
    if align is not None:
        paragraph.alignment = align


def _add_run(paragraph, text, *, bold=False, size=10.5, color=None, font_name='Arial'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _set_bottom_border_only(cell):
    _set_cell_border(
        cell,
        top={'val': 'nil'},
        left={'val': 'nil'},
        right={'val': 'nil'},
        bottom={'val': 'single', 'sz': '10', 'space': '0', 'color': '000000'},
    )


def _clear_cell(cell):
    cell._tc.clear_content()
    cell._tc.append(OxmlElement('w:p'))


def _value_line(parent_cell, placeholder, *, font_size=10.5):
    line_table = parent_cell.add_table(rows=1, cols=1)
    line_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    line_table.autofit = False
    _make_table_borderless(line_table)
    value_cell = line_table.cell(0, 0)
    _set_cell_margins(value_cell, top=30, start=0, bottom=10, end=0)
    _set_bottom_border_only(value_cell)
    p = value_cell.paragraphs[0]
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, placeholder, size=font_size)
    return value_cell


def _line_value_cell(cell, placeholder, *, font_size=10.5, align=None):
    _clear_cell(cell)
    _set_cell_margins(cell, top=30, start=20, bottom=10, end=20)
    _set_bottom_border_only(cell)
    p = cell.paragraphs[0]
    _set_paragraph_format(p, before=0, after=0, align=align)
    _add_run(p, placeholder, size=font_size)
    return cell


def _label_value_box(cell, rows, *, col_widths=(Inches(0.95), Inches(1.2)), font_size=10):
    box = cell.add_table(rows=len(rows), cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.autofit = False
    box.style = 'Table Grid'
    _set_table_borders(box, size='10')
    for row in box.rows:
        _set_col_widths(row, col_widths)
        for item in row.cells:
            item.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(item, top=40, start=60, bottom=40, end=60)

    for idx, (label, placeholder) in enumerate(rows):
        p = box.cell(idx, 0).paragraphs[0]
        _set_paragraph_format(p, before=0, after=0)
        _add_run(p, label, bold=True, size=font_size)
        _clear_cell(box.cell(idx, 1))
        _value_line(box.cell(idx, 1), placeholder, font_size=font_size)
    return box


def _party_block(cell, heading, placeholders):
    table = cell.add_table(rows=1 + len(placeholders), cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    _make_table_borderless(table)
    heading_cell = table.cell(0, 0)
    _set_cell_margins(heading_cell, top=0, start=0, bottom=20, end=0)
    hp = heading_cell.paragraphs[0]
    _set_paragraph_format(hp, before=0, after=0)
    _add_run(hp, heading, bold=True, size=12)
    for idx, placeholder in enumerate(placeholders, start=1):
        line_cell = table.cell(idx, 0)
        _set_cell_margins(line_cell, top=20, start=0, bottom=10, end=0)
        _set_bottom_border_only(line_cell)
        p = line_cell.paragraphs[0]
        _set_paragraph_format(p, before=0, after=0)
        _add_run(p, placeholder, size=10)
    return table


def _owners_risk_box(cell):
    box = cell.add_table(rows=3, cols=4)
    box.alignment = WD_TABLE_ALIGNMENT.LEFT
    box.autofit = False
    box.style = 'Table Grid'
    _set_table_borders(box, size='10')

    widths = [Inches(0.95), Inches(1.45), Inches(0.65), Inches(1.45)]
    for row in box.rows:
        _set_col_widths(row, widths)
        for item in row.cells:
            item.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _set_cell_margins(item, top=35, start=45, bottom=35, end=45)

    p = box.cell(0, 0).paragraphs[0]
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, 'Company', bold=True, size=10)
    merged = box.cell(0, 1).merge(box.cell(0, 3))
    _line_value_cell(merged, '{{INSURANCE_COMPANY}}', font_size=10)

    row_two = [
        ('Policy No.', '{{INSURANCE_POLICY_NO}}'),
        ('Date', '{{INSURANCE_DATE}}'),
    ]
    row_three = [
        ('Amount', '{{INSURANCE_AMOUNT}}'),
        ('Risk', '{{INSURANCE_RISK}}'),
    ]

    for row_idx, row_data in enumerate((row_two, row_three), start=1):
        for pair_idx, (label, placeholder) in enumerate(row_data):
            label_cell = box.cell(row_idx, pair_idx * 2)
            value_cell = box.cell(row_idx, pair_idx * 2 + 1)
            p = label_cell.paragraphs[0]
            _set_paragraph_format(p, before=0, after=0)
            _add_run(p, label, bold=True, size=10)
            _line_value_cell(value_cell, placeholder, font_size=10)

    return box


def build_lr_template(output_path=None):
    output_path = Path(output_path) if output_path else TEMPLATE_PATH

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.3)
    section.right_margin = Inches(0.3)

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)

    header = doc.add_table(rows=1, cols=3)
    header.autofit = False
    header.alignment = WD_TABLE_ALIGNMENT.LEFT
    _make_table_borderless(header)
    _set_col_widths(header.rows[0], [Inches(0.8), Inches(5.9), Inches(0.8)])

    logo_cell = header.cell(0, 0)
    _set_cell_margins(logo_cell, top=0, start=0, bottom=0, end=0)
    if LOGO_PATH.exists():
        logo_run = logo_cell.paragraphs[0].add_run()
        logo_run.add_picture(str(LOGO_PATH), width=Inches(0.55))

    title_cell = header.cell(0, 1)
    _set_cell_margins(title_cell, top=0, start=40, bottom=0, end=40)
    p = title_cell.paragraphs[0]
    _set_paragraph_format(p, before=0, after=1)
    _add_run(p, 'Godamwale Trading & Logistics Pvt Ltd', bold=True, size=24)
    p = title_cell.add_paragraph()
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, '711, Swastik Chambers, Sion Trombay Road, Chembur, Mumbai, Maharashtra - 400071', size=11)
    p = title_cell.add_paragraph()
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, 'Email: info@godamwale.com, Website: www.godamwale.com', size=11)

    contact_cell = header.cell(0, 2)
    _set_cell_margins(contact_cell, top=0, start=0, bottom=0, end=0)
    p = contact_cell.paragraphs[0]
    _set_paragraph_format(p, before=0, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _add_run(p, 'M: 8291540681', size=10)

    doc.add_paragraph()

    body = doc.add_table(rows=1, cols=2)
    body.autofit = False
    body.alignment = WD_TABLE_ALIGNMENT.LEFT
    _make_table_borderless(body)
    _set_col_widths(body.rows[0], [Inches(6.1), Inches(4.0)])

    left_panel = body.cell(0, 0)
    right_panel = body.cell(0, 1)
    _set_cell_margins(left_panel, top=0, start=0, bottom=0, end=20)
    _set_cell_margins(right_panel, top=0, start=20, bottom=0, end=0)

    top_row = left_panel.add_table(rows=1, cols=3)
    top_row.autofit = False
    top_row.alignment = WD_TABLE_ALIGNMENT.LEFT
    _make_table_borderless(top_row)
    _set_col_widths(top_row.rows[0], [Inches(1.45), Inches(2.45), Inches(2.0)])

    copy_cell = top_row.cell(0, 0)
    _set_cell_margins(copy_cell, top=0, start=0, bottom=0, end=15)
    badge = copy_cell.add_table(rows=1, cols=1)
    badge.autofit = False
    badge.style = 'Table Grid'
    _set_table_borders(badge, size='12')
    bp = badge.cell(0, 0).paragraphs[0]
    _set_paragraph_format(bp, before=0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(bp, 'CONSIGNEE COPY', bold=True, size=11.5, color='C62828')

    delivery_cell = top_row.cell(0, 1)
    _set_cell_margins(delivery_cell, top=0, start=10, bottom=0, end=20)
    p = delivery_cell.paragraphs[0]
    _set_paragraph_format(p, before=0, after=2)
    _add_run(p, 'Address of Delivery Office', bold=True, size=11)
    _value_line(delivery_cell, '{{DELIVERY_OFFICE_LINE_1}}', font_size=9.3)
    _value_line(delivery_cell, '{{DELIVERY_OFFICE_LINE_2}}', font_size=9.3)
    _value_line(delivery_cell, '{{DELIVERY_OFFICE_LINE_3}}', font_size=9.3)

    note_cell = top_row.cell(0, 2)
    _set_cell_margins(note_cell, top=0, start=10, bottom=0, end=0)
    note = note_cell.add_table(rows=5, cols=2)
    note.autofit = False
    note.style = 'Table Grid'
    _set_table_borders(note, size='10')
    note.cell(0, 0).merge(note.cell(0, 1))
    for row in note.rows:
        _set_col_widths(row, [Inches(0.72), Inches(1.18)])
        for c in row.cells:
            _set_cell_margins(c, top=35, start=45, bottom=35, end=45)
    p = note.cell(0, 0).paragraphs[0]
    _set_paragraph_format(p, before=0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    _add_run(p, 'CONSIGNMENT NOTE', bold=True, size=11.5)
    for row_idx, (label, placeholder) in enumerate([
        ('No.', '{{LR_NUMBER}}'),
        ('Date', '{{LR_DATE}}'),
        ('From', '{{FROM_LOCATION}}'),
        ('To', '{{TO_LOCATION}}'),
    ], start=1):
        lp = note.cell(row_idx, 0).paragraphs[0]
        _set_paragraph_format(lp, before=0, after=0)
        _add_run(lp, label, bold=True, size=10)
        _line_value_cell(note.cell(row_idx, 1), placeholder, font_size=10)

    p = left_panel.add_paragraph()
    _set_paragraph_format(p, before=7, after=3)
    _add_run(p, 'AT OWNERS RISK', bold=True, size=11)
    _owners_risk_box(left_panel)

    p = left_panel.add_paragraph()
    _set_paragraph_format(p, before=5, after=0)
    parties = left_panel.add_table(rows=1, cols=2)
    parties.autofit = False
    parties.alignment = WD_TABLE_ALIGNMENT.LEFT
    _make_table_borderless(parties)
    _set_col_widths(parties.rows[0], [Inches(3.0), Inches(3.0)])
    _set_cell_margins(parties.cell(0, 0), top=0, start=0, bottom=0, end=15)
    _set_cell_margins(parties.cell(0, 1), top=0, start=15, bottom=0, end=0)
    _party_block(parties.cell(0, 0), "Consignor's Name & Address", [
        '{{CONSIGNOR_LINE_1}}', '{{CONSIGNOR_LINE_2}}', '{{CONSIGNOR_LINE_3}}', '{{CONSIGNOR_LINE_4}}',
        '{{CONSIGNOR_LINE_5}}',
    ])
    _party_block(parties.cell(0, 1), "Consignee's Name & Address", [
        '{{CONSIGNEE_LINE_1}}', '{{CONSIGNEE_LINE_2}}', '{{CONSIGNEE_LINE_3}}', '{{CONSIGNEE_LINE_4}}',
        '{{CONSIGNEE_LINE_5}}',
    ])

    info = right_panel.add_table(rows=10, cols=2)
    info.autofit = False
    info.style = 'Table Grid'
    _set_table_borders(info, size='10')
    for row in info.rows:
        _set_col_widths(row, [Inches(1.45), Inches(2.45)])
        for c in row.cells:
            _set_cell_margins(c, top=35, start=45, bottom=35, end=45)
            c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    simple_rows = [
        ('PAN', '{{PAN_NUMBER}}'),
        ('GSTIN', '{{GSTIN_NUMBER}}'),
        ('Veh no.', '{{VEHICLE_NO}}'),
        ('Mode of Packing', '{{MODE_OF_PACKING}}'),
        ('Invoice no.', '{{INVOICE_NO}}'),
        ("Consignor GST no.", '{{CONSIGNOR_GST_NO}}'),
        ("Consignee GST no.", '{{CONSIGNEE_GST_NO}}'),
        ('Vehicle Type', '{{VEHICLE_TYPE}}'),
        ('Remarks', '{{REMARKS}}'),
    ]
    simple_row_positions = [0, 1, 2, 4, 5, 6, 7, 8, 9]
    for pos, (label, placeholder) in zip(simple_row_positions, simple_rows):
        p = info.cell(pos, 0).paragraphs[0]
        _set_paragraph_format(p, before=0, after=0)
        _add_run(p, label, bold=True, size=9.8)
        _line_value_cell(info.cell(pos, 1), placeholder, font_size=9.8)

    p = info.cell(3, 0).paragraphs[0]
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, 'GST to be Paid by', bold=True, size=9.8)
    gst_cell = info.cell(3, 1)
    _clear_cell(gst_cell)
    for idx, (box_placeholder, label) in enumerate([
        ('{{GST_CONSIGNEE_BOX}}', 'Consignee'),
        ('{{GST_CONSIGNOR_BOX}}', 'Consignor'),
        ('{{GST_TRANSPORTER_BOX}}', 'Transporter'),
    ]):
        gp = gst_cell.paragraphs[0] if idx == 0 else gst_cell.add_paragraph()
        _set_paragraph_format(gp, before=0, after=0)
        _add_run(gp, box_placeholder, font_name='Segoe UI Symbol', size=11)
        _add_run(gp, f' {label}', size=9.8)

    doc.add_paragraph()

    items = doc.add_table(rows=2, cols=5)
    items.autofit = False
    items.style = 'Table Grid'
    _set_table_borders(items, size='10')
    widths = [Inches(1.0), Inches(3.7), Inches(1.25), Inches(1.25), Inches(2.8)]
    for row in items.rows:
        _set_col_widths(row, widths)
        for c in row.cells:
            _set_cell_margins(c, top=40, start=50, bottom=40, end=50)
    for cell, text in zip(items.rows[0].cells, ['Packages', 'Description', 'Actual Weight', 'Charged Weight', 'Amount to Pay/Paid']):
        p = cell.paragraphs[0]
        _set_paragraph_format(p, before=0, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
        _add_run(p, text, bold=True, size=10.5)
    for cell, placeholder in zip(items.rows[1].cells, ['{{ITEM_PACKAGES}}', '{{ITEM_DESCRIPTION}}', '{{ITEM_ACTUAL_WEIGHT}}', '{{ITEM_CHARGED_WEIGHT}}', '{{ITEM_AMOUNT}}']):
        p = cell.paragraphs[0]
        _set_paragraph_format(p, before=0, after=0)
        _add_run(p, placeholder, size=9.5)

    doc.add_paragraph()

    bottom = doc.add_table(rows=1, cols=4)
    bottom.autofit = False
    bottom.alignment = WD_TABLE_ALIGNMENT.LEFT
    _make_table_borderless(bottom)
    _set_col_widths(bottom.rows[0], [Inches(0.8), Inches(3.8), Inches(1.3), Inches(4.3)])

    p = bottom.cell(0, 0).paragraphs[0]
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, 'Value', bold=True, size=11)
    _line_value_cell(bottom.cell(0, 1), '{{VALUE}}', font_size=11)

    p = bottom.cell(0, 2).paragraphs[0]
    _set_paragraph_format(p, before=0, after=0)
    _add_run(p, "Consignor's sign", bold=True, size=11)
    _line_value_cell(bottom.cell(0, 3), '', font_size=11)

    doc.save(output_path)
    return output_path


if __name__ == '__main__':
    build_lr_template()
