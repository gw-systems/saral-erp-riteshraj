
import copy
import logging
import os
import re
import tempfile
import textwrap
from pathlib import Path

from docx import Document

logger = logging.getLogger(__name__)

ROOT = Path(__file__).resolve().parents[2]
TEMPLATE_PATH = ROOT / 'LR.docx'
LOGO_TEMPLATE_PATH = ROOT / 'static' / 'images' / 'GW Logo.png'
PAN_NUMBER = 'AAGCG3326A'
GSTIN_NUMBER = '27AAGCG3326A1Z3'
PLACEHOLDER_RE = re.compile(r'{{[A-Z0-9_]+}}')
ITEM_PLACEHOLDERS = {
    '{{ITEM_PACKAGES}}': 'packages',
    '{{ITEM_DESCRIPTION}}': 'description',
    '{{ITEM_ACTUAL_WEIGHT}}': 'actual_weight',
    '{{ITEM_CHARGED_WEIGHT}}': 'charged_weight',
    '{{ITEM_AMOUNT}}': 'amount',
}
REQUIRED_TEMPLATE_PLACEHOLDERS = {
    '{{LR_NUMBER}}',
    '{{LR_DATE}}',
    '{{FROM_LOCATION}}',
    '{{TO_LOCATION}}',
    '{{PAN_NUMBER}}',
    '{{GSTIN_NUMBER}}',
    '{{VEHICLE_NO}}',
    '{{MODE_OF_PACKING}}',
    '{{INVOICE_NO}}',
    '{{CONSIGNOR_GST_NO}}',
    '{{CONSIGNEE_GST_NO}}',
    '{{VEHICLE_TYPE}}',
    '{{REMARKS}}',
    '{{INSURANCE_COMPANY}}',
    '{{INSURANCE_POLICY_NO}}',
    '{{INSURANCE_DATE}}',
    '{{INSURANCE_AMOUNT}}',
    '{{INSURANCE_RISK}}',
    '{{VALUE}}',
    '{{DELIVERY_OFFICE_LINE_1}}',
    '{{DELIVERY_OFFICE_LINE_2}}',
    '{{DELIVERY_OFFICE_LINE_3}}',
    '{{CONSIGNOR_LINE_1}}',
    '{{CONSIGNOR_LINE_2}}',
    '{{CONSIGNOR_LINE_3}}',
    '{{CONSIGNEE_LINE_1}}',
    '{{CONSIGNEE_LINE_2}}',
    '{{CONSIGNEE_LINE_3}}',
    '{{GST_CONSIGNEE_BOX}}',
    '{{GST_CONSIGNOR_BOX}}',
    '{{GST_TRANSPORTER_BOX}}',
    *ITEM_PLACEHOLDERS.keys(),
}


def _ensure_template_exists():
    if TEMPLATE_PATH.exists():
        return TEMPLATE_PATH
    from .lr_template_builder import build_lr_template
    return build_lr_template(TEMPLATE_PATH)


def _format_date(value):
    if not value:
        return ''
    try:
        return value.strftime('%d/%m/%Y')
    except AttributeError:
        return str(value)


def _split_text_lines(*parts, width, max_lines):
    lines = []
    for part in parts:
        if not part:
            continue
        normalized = str(part).replace('\r', '\n')
        for chunk in normalized.split('\n'):
            chunk = chunk.strip()
            if not chunk:
                continue
            wrapped = textwrap.wrap(chunk, width=width, break_long_words=False, break_on_hyphens=False)
            lines.extend(wrapped or [''])
    if not lines:
        return [''] * max_lines
    if len(lines) > max_lines:
        collapsed = lines[: max_lines - 1]
        collapsed.append(' '.join(lines[max_lines - 1:]))
        lines = collapsed
    return lines + [''] * (max_lines - len(lines))


def _gst_boxes(selected_value):
    return {
        '{{GST_CONSIGNEE_BOX}}': '\u2611' if selected_value == 'consignee' else '\u2610',
        '{{GST_CONSIGNOR_BOX}}': '\u2611' if selected_value == 'consignor' else '\u2610',
        '{{GST_TRANSPORTER_BOX}}': '\u2611' if selected_value == 'transporter' else '\u2610',
    }


def _build_placeholder_map(lr):
    data = {
        '{{LR_NUMBER}}': lr.lr_number or '',
        '{{LR_DATE}}': _format_date(lr.lr_date),
        '{{FROM_LOCATION}}': lr.from_location or '',
        '{{TO_LOCATION}}': lr.to_location or '',
        '{{PAN_NUMBER}}': PAN_NUMBER,
        '{{GSTIN_NUMBER}}': GSTIN_NUMBER,
        '{{VEHICLE_NO}}': lr.vehicle_no or '',
        '{{MODE_OF_PACKING}}': lr.mode_of_packing or '',
        '{{INVOICE_NO}}': lr.invoice_no or '',
        '{{CONSIGNOR_GST_NO}}': lr.consignor_gst_no or '',
        '{{CONSIGNEE_GST_NO}}': lr.consignee_gst_no or '',
        '{{VEHICLE_TYPE}}': lr.vehicle_type or '',
        '{{REMARKS}}': lr.remarks or '',
        '{{INSURANCE_COMPANY}}': lr.insurance_company or '',
        '{{INSURANCE_POLICY_NO}}': lr.insurance_policy_no or '',
        '{{INSURANCE_DATE}}': lr.insurance_date or '',
        '{{INSURANCE_AMOUNT}}': lr.insurance_amount or '',
        '{{INSURANCE_RISK}}': lr.insurance_risk or '',
        '{{VALUE}}': lr.value or '',
    }

    for idx, line in enumerate(_split_text_lines(lr.delivery_office_address, width=32, max_lines=3), start=1):
        data[f'{{{{DELIVERY_OFFICE_LINE_{idx}}}}}'] = line

    for idx, line in enumerate(_split_text_lines(lr.consignor_name, lr.consignor_address, width=33, max_lines=3), start=1):
        data[f'{{{{CONSIGNOR_LINE_{idx}}}}}'] = line

    for idx, line in enumerate(_split_text_lines(lr.consignee_name, lr.consignee_address, width=33, max_lines=3), start=1):
        data[f'{{{{CONSIGNEE_LINE_{idx}}}}}'] = line

    data.update(_gst_boxes(lr.gst_paid_by))
    return data


def _iter_paragraphs(parent):
    if hasattr(parent, 'paragraphs'):
        for paragraph in parent.paragraphs:
            yield paragraph
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)


def _template_placeholders(document):
    matches = set()
    for paragraph in _iter_paragraphs(document):
        for match in PLACEHOLDER_RE.findall(paragraph.text):
            matches.add(match)
    return matches


def _replace_placeholders_in_paragraph(paragraph, mapping):
    for run in paragraph.runs:
        text = run.text
        if not text:
            continue
        updated = text
        for placeholder, value in mapping.items():
            if placeholder in updated:
                updated = updated.replace(placeholder, value)
        if updated != text:
            run.text = updated


def _find_item_table(document):
    for table in document.tables:
        for row in table.rows:
            if any(token in cell.text for token in ITEM_PLACEHOLDERS for cell in row.cells):
                return table, row
    raise RuntimeError('LR template item row placeholders are missing.')


def _fill_item_row(row, item):
    values = {
        '{{ITEM_PACKAGES}}': getattr(item, 'packages', '') or '',
        '{{ITEM_DESCRIPTION}}': getattr(item, 'description', '') or '',
        '{{ITEM_ACTUAL_WEIGHT}}': getattr(item, 'actual_weight', '') or '',
        '{{ITEM_CHARGED_WEIGHT}}': getattr(item, 'charged_weight', '') or '',
        '{{ITEM_AMOUNT}}': getattr(item, 'amount', '') or '',
    }
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            _replace_placeholders_in_paragraph(paragraph, values)


def _fill_line_items(document, line_items):
    table, template_row = _find_item_table(document)
    items = list(line_items or [])
    if not items:
        items = [type('BlankItem', (), {field: '' for field in ITEM_PLACEHOLDERS.values()})()]

    template_row_xml = copy.deepcopy(template_row._tr)
    _fill_item_row(template_row, items[0])
    for item in items[1:]:
        new_tr = copy.deepcopy(template_row_xml)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        _fill_item_row(new_row, item)


def _unresolved_placeholders(document):
    return sorted(_template_placeholders(document))


def _validate_template(document):
    present = _template_placeholders(document)
    missing = sorted(REQUIRED_TEMPLATE_PLACEHOLDERS - present)
    if missing:
        raise RuntimeError(
            'LR template is missing required placeholders: ' + ', '.join(missing)
        )


def generate_lr_docx(lr, line_items):
    template_path = _ensure_template_exists()
    if not template_path.exists():
        raise FileNotFoundError(f'LR template not found at {template_path}.')

    document = Document(str(template_path))
    _validate_template(document)
    mapping = _build_placeholder_map(lr)

    for paragraph in _iter_paragraphs(document):
        _replace_placeholders_in_paragraph(paragraph, mapping)

    _fill_line_items(document, line_items)

    unresolved = _unresolved_placeholders(document)
    if unresolved:
        raise RuntimeError(
            'LR template has unresolved placeholders after generation: ' + ', '.join(unresolved)
        )

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    document.save(tmp.name)
    logger.info('LR DOCX generated: %s', tmp.name)
    return tmp.name


def _find_soffice():
    import shutil

    env_path = os.environ.get('SOFFICE_PATH')
    if env_path and os.path.isfile(env_path):
        return env_path

    for binary in ('soffice', 'soffice.exe', 'libreoffice', 'libreoffice.exe'):
        path = shutil.which(binary)
        if path:
            return path

    windows_candidates = [
        os.path.join(os.environ.get('PROGRAMFILES', ''), 'LibreOffice', 'program', 'soffice.com'),
        os.path.join(os.environ.get('PROGRAMFILES', ''), 'LibreOffice', 'program', 'soffice.exe'),
        os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'LibreOffice', 'program', 'soffice.com'),
        os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'LibreOffice', 'program', 'soffice.exe'),
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Programs', 'LibreOffice', 'program', 'soffice.com'),
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'Programs', 'LibreOffice', 'program', 'soffice.exe'),
        r'C:\Program Files\LibreOffice\program\soffice.com',
        r'C:\Program Files\LibreOffice\program\soffice.exe',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.com',
        r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
    ]
    unix_candidates = [
        '/Applications/LibreOffice.app/Contents/MacOS/soffice',
        '/usr/local/bin/soffice',
        '/usr/bin/soffice',
    ]

    for candidate in windows_candidates + unix_candidates:
        if os.path.isfile(candidate):
            return candidate
    return None


def _find_poppler_bin_dir():
    import shutil

    env_path = os.environ.get('POPPLER_PATH')
    if env_path and os.path.isdir(env_path):
        pdftoppm = os.path.join(env_path, 'pdftoppm.exe' if os.name == 'nt' else 'pdftoppm')
        if os.path.isfile(pdftoppm):
            return env_path

    direct_binary = shutil.which('pdftoppm.exe' if os.name == 'nt' else 'pdftoppm')
    if direct_binary:
        return os.path.dirname(direct_binary)

    local_candidates = [
        ROOT / 'tools' / 'poppler',
        ROOT / 'poppler',
    ]
    for base in local_candidates:
        if not base.exists():
            continue
        matches = sorted(base.glob('**/pdftoppm.exe' if os.name == 'nt' else '**/pdftoppm'))
        if matches:
            return str(matches[0].parent)

    windows_candidates = [
        Path(r'C:\Program Files\poppler\Library\bin'),
        Path(r'C:\ProgramData\chocolatey\lib\poppler\tools\poppler-25.12.0\Library\bin'),
    ]
    for candidate in windows_candidates:
        binary = candidate / ('pdftoppm.exe' if os.name == 'nt' else 'pdftoppm')
        if binary.is_file():
            return str(candidate)

    return None


def _convert_docx_to_pdf_soffice(docx_path, pdf_path):
    import subprocess
    import time

    soffice = _find_soffice()
    if not soffice:
        raise FileNotFoundError('LibreOffice not found')

    out_dir = os.path.dirname(pdf_path) or tempfile.gettempdir()
    env = os.environ.copy()
    soffice_dir = os.path.dirname(soffice)
    env['PATH'] = soffice_dir + os.pathsep + env.get('PATH', '')
    result = subprocess.run(
        [soffice, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, docx_path],
        capture_output=True,
        text=True,
        timeout=120,
        env=env,
    )
    expected = os.path.join(out_dir, os.path.basename(docx_path).replace('.docx', '.pdf'))
    for _ in range(20):
        if os.path.exists(expected):
            if expected != pdf_path:
                os.replace(expected, pdf_path)
            return pdf_path
        time.sleep(0.5)
    raise RuntimeError(
        'LibreOffice conversion produced no output.\n'
        f'returncode: {result.returncode}\nstdout: {result.stdout}\nstderr: {result.stderr}'
    )


def _convert_docx_to_pdf_word(docx_path, pdf_path):
    import subprocess

    if os.name != 'nt':
        raise FileNotFoundError('Microsoft Word conversion is only available on Windows')

    escaped_docx = os.path.abspath(docx_path).replace("'", "''")
    escaped_pdf = os.path.abspath(pdf_path).replace("'", "''")
    script = (
        "$ErrorActionPreference='Stop';"
        "$word=$null;$doc=$null;"
        f"$docx='{escaped_docx}';"
        f"$pdf='{escaped_pdf}';"
        "try {"
        "  $word = New-Object -ComObject Word.Application;"
        "  $word.Visible = $false;"
        "  $word.DisplayAlerts = 0;"
        "  $doc = $word.Documents.Open($docx, $false, $true);"
        "  $doc.ExportAsFixedFormat($pdf, 17);"
        "} finally {"
        "  if ($doc -ne $null) { $doc.Close([ref]$false) }"
        "  if ($word -ne $null) { $word.Quit() }"
        "}"
    )
    result = subprocess.run(
        ['powershell', '-NoProfile', '-NonInteractive', '-Command', script],
        capture_output=True,
        text=True,
        timeout=120,
    )
    if os.path.exists(pdf_path):
        return pdf_path
    raise RuntimeError(
        'Word conversion produced no output.\n'
        f'returncode: {result.returncode}\nstdout: {result.stdout}\nstderr: {result.stderr}'
    )


def generate_lr_pdf(lr, line_items):
    docx_path = generate_lr_docx(lr, line_items)
    pdf_path = docx_path.replace('.docx', '.pdf')
    try:
        _convert_docx_to_pdf_soffice(docx_path, pdf_path)
        logger.info('LR PDF generated (LibreOffice): %s', pdf_path)
        return pdf_path
    except (FileNotFoundError, Exception) as libre_err:
        logger.warning('LibreOffice failed: %s, trying Word/docx2pdf fallback', libre_err)
        try:
            _convert_docx_to_pdf_word(docx_path, pdf_path)
            logger.info('LR PDF generated (Word): %s', pdf_path)
            return pdf_path
        except (FileNotFoundError, Exception) as word_err:
            logger.warning('Word fallback failed: %s, trying docx2pdf', word_err)
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
                logger.info('LR PDF generated (docx2pdf): %s', pdf_path)
                return pdf_path
            except ImportError:
                raise RuntimeError(
                    'PDF conversion requires LibreOffice, Microsoft Word, or docx2pdf.\n'
                    'Windows local fallback can use Microsoft Word automation if installed.\n'
                    'Or install docx2pdf: pip install docx2pdf'
                )
            except SystemExit:
                raise RuntimeError(
                    'docx2pdf failed (Word automation error). Install LibreOffice or use the Microsoft Word fallback instead.'
                )
            except Exception as exc:
                raise RuntimeError(f'PDF conversion failed: {exc}')
    finally:
        try:
            os.unlink(docx_path)
        except OSError:
            pass


def generate_lr_image(lr, line_items):
    pdf_path = generate_lr_pdf(lr, line_items)
    try:
        from pdf2image import convert_from_path
        poppler_path = _find_poppler_bin_dir()
        convert_kwargs = {
            'dpi': 200,
            'first_page': 1,
            'last_page': 1,
        }
        if poppler_path:
            convert_kwargs['poppler_path'] = poppler_path
        images = convert_from_path(pdf_path, **convert_kwargs)
        if not images:
            raise RuntimeError('pdf2image returned no images')
        image_path = pdf_path.replace('.pdf', '.png')
        images[0].save(image_path, 'PNG')
        logger.info('LR image generated: %s', image_path)
        return image_path
    except ImportError:
        raise RuntimeError(
            'pdf2image is not installed. Install it with: pip install pdf2image\n'
            'Also requires poppler in the local environment.'
        )
    except Exception as exc:
        raise RuntimeError(f'Image generation failed: {exc}')
    finally:
        try:
            os.unlink(pdf_path)
        except OSError:
            pass
