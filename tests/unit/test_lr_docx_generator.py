import os
import re
from datetime import date

import pytest
from docx import Document

from operations.models_lr import LRLineItem, LorryReceipt
from operations.services.lr_docx_generator import generate_lr_docx, generate_lr_pdf


def _walk_paragraph_text(parent):
    if hasattr(parent, 'paragraphs'):
        for paragraph in parent.paragraphs:
            text = paragraph.text.strip()
            if text:
                yield text
    if hasattr(parent, 'tables'):
        for table in parent.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _walk_paragraph_text(cell)


@pytest.fixture
def lr_sample(db, admin_user, test_project):
    lr = LorryReceipt.objects.create(
        lr_date=date(2026, 3, 14),
        project=test_project,
        from_location='Mumbai',
        to_location='Pune',
        vehicle_no='MH12AB1234',
        vehicle_type='32 FT OPEN',
        delivery_office_address='Godamwale Delivery Hub\nGate 3, Bhosari MIDC, Pune',
        consignor_name='Dummy Consignor Pvt Ltd',
        consignor_address='Plot 18, TTC Industrial Area, Navi Mumbai',
        consignee_name='Dummy Consignee Industries',
        consignee_address='Warehouse 4, Chakan Phase 2, Pune',
        consignor_gst_no='27AAAAA0000A1Z5',
        consignee_gst_no='27BBBBB0000B1Z6',
        invoice_no='INV-42',
        gst_paid_by='transporter',
        mode_of_packing='Loose Cartons',
        value='250000',
        remarks='Handle with care',
        insurance_company='ICICI Lombard',
        insurance_policy_no='PL-9981',
        insurance_date='14/03/2026',
        insurance_amount='250000',
        insurance_risk='Transit Risk',
        created_by=admin_user,
        last_modified_by=admin_user,
    )
    LRLineItem.objects.create(
        lr=lr,
        packages='12',
        description='Cartons of garments',
        actual_weight='980 KG',
        charged_weight='1000 KG',
        amount='12500',
        order=1,
    )
    return lr


@pytest.mark.django_db
def test_generate_lr_docx_replaces_all_placeholders(lr_sample):
    docx_path = generate_lr_docx(lr_sample, list(lr_sample.line_items.all()))
    try:
        document = Document(docx_path)
        text = '\n'.join(_walk_paragraph_text(document))

        assert not re.findall(r'{{[A-Z0-9_]+}}', text)
        assert "Consignor's Name & Address" in text
        assert "Consignee's Name & Address" in text
        assert 'Dummy Consignor Pvt Ltd' in text
        assert 'Dummy Consignee Industries' in text
        assert 'Godamwale Delivery Hub' in text
        assert 'ICICI Lombard' in text
        assert '250000' in text
        assert 'Transit Risk' in text
        assert '\u2611 Transporter' in text
        assert '\u2610 Consignee' in text
        assert '\u2610 Consignor' in text
        assert 'Cartons of garments' in text
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)


@pytest.mark.django_db
def test_generate_lr_docx_keeps_distinct_line_items(lr_sample):
    LRLineItem.objects.filter(lr=lr_sample).delete()
    LRLineItem.objects.create(
        lr=lr_sample,
        packages='1',
        description='Good Goods',
        actual_weight='100 KG',
        charged_weight='100 KG',
        amount='1000',
        order=1,
    )
    LRLineItem.objects.create(
        lr=lr_sample,
        packages='1',
        description='Bad Goods',
        actual_weight='80 KG',
        charged_weight='80 KG',
        amount='500',
        order=2,
    )
    LRLineItem.objects.create(
        lr=lr_sample,
        packages='1',
        description='Avg Goods',
        actual_weight='90 KG',
        charged_weight='90 KG',
        amount='700',
        order=3,
    )

    docx_path = generate_lr_docx(lr_sample, list(lr_sample.line_items.all()))
    try:
        document = Document(docx_path)
        text = '\n'.join(_walk_paragraph_text(document))
        assert text.count('Good Goods') == 1
        assert text.count('Bad Goods') == 1
        assert text.count('Avg Goods') == 1
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)


@pytest.mark.django_db
def test_generate_lr_pdf_creates_pdf_when_converter_is_available(lr_sample):
    try:
        pdf_path = generate_lr_pdf(lr_sample, list(lr_sample.line_items.all()))
    except RuntimeError as exc:
        pytest.skip(f'No local PDF converter available for LR test: {exc}')

    try:
        with open(pdf_path, 'rb') as handle:
            header = handle.read(4)
        assert header == b'%PDF'
    finally:
        if os.path.exists(pdf_path):
            os.unlink(pdf_path)
